#!/usr/bin/env python3
"""
Report Generator Module
CSV and DOCX report generation
"""

import csv
import os
from datetime import datetime
from typing import List, Dict, Any
from dataclasses import dataclass

# DOCX generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class ReportConfig:
    report_type: str = "weekly"  # weekly or monthly
    company_name: str = "Company"
    team_name: str = "Infrastructure Team"
    output_dir: str = "./output"


class ReportGenerator:
    def __init__(self, config: ReportConfig = None):
        self.config = config or ReportConfig()
        os.makedirs(self.config.output_dir, exist_ok=True)
        
    def _get_report_title(self) -> str:
        now = datetime.now()
        if self.config.report_type == "weekly":
            week_num = now.isocalendar()[1]
            return f"{now.year} Week {week_num} Infrastructure Health Check Report"
        else:
            return f"{now.year} {now.strftime('%B')} Infrastructure Health Check Report"
    
    def _get_filename_prefix(self) -> str:
        now = datetime.now()
        if self.config.report_type == "weekly":
            week_num = now.isocalendar()[1]
            return f"infra_check_{now.year}_W{week_num:02d}"
        else:
            return f"infra_check_{now.year}_{now.month:02d}"
    
    def generate_csv(self, results: List[Dict], summary: Dict) -> str:
        """Generate CSV report"""
        filename = f"{self._get_filename_prefix()}.csv"
        filepath = os.path.join(self.config.output_dir, filename)
        
        with open(filepath, 'w', newline='', encoding='utf-8-sig') as f:
            # Header information
            f.write(f"# {self._get_report_title()}\n")
            f.write(f"# Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"# Company: {self.config.company_name}\n")
            f.write(f"# Team: {self.config.team_name}\n")
            f.write(f"# Total Items: {summary.get('total', 0)}\n")
            f.write(f"# OK: {summary.get('ok', 0)} / Warning: {summary.get('warning', 0)} / Critical: {summary.get('critical', 0)} / Unknown: {summary.get('unknown', 0)}\n")
            f.write("\n")
            
            # Data
            if results:
                writer = csv.DictWriter(f, fieldnames=results[0].keys())
                writer.writeheader()
                writer.writerows(results)
        
        return filepath
    
    def generate_docx(self, results: List[Dict], summary: Dict) -> str:
        """Generate DOCX report"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx library is not installed. Run: pip install python-docx")
        
        filename = f"{self._get_filename_prefix()}.docx"
        filepath = os.path.join(self.config.output_dir, filename)
        
        doc = Document()
        
        # Title
        title = doc.add_heading(self._get_report_title(), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Report info
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n").bold = False
        info_para.add_run(f"Company: {self.config.company_name} | Team: {self.config.team_name}")
        
        doc.add_paragraph()
        
        # Summary section
        doc.add_heading('1. Executive Summary', level=1)
        
        summary_table = doc.add_table(rows=2, cols=5)
        summary_table.style = 'Table Grid'
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        hdr_cells = summary_table.rows[0].cells
        headers = ['Total Items', 'OK', 'Warning', 'Critical', 'Unknown']
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in hdr_cells[i].paragraphs[0].runs:
                run.bold = True
        
        # Data
        data_cells = summary_table.rows[1].cells
        data = [
            str(summary.get('total', 0)),
            str(summary.get('ok', 0)),
            str(summary.get('warning', 0)),
            str(summary.get('critical', 0)),
            str(summary.get('unknown', 0))
        ]
        colors = [None, RGBColor(0, 128, 0), RGBColor(255, 165, 0), RGBColor(255, 0, 0), RGBColor(128, 128, 128)]
        
        for i, (value, color) in enumerate(zip(data, colors)):
            data_cells[i].text = value
            data_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if color:
                for run in data_cells[i].paragraphs[0].runs:
                    run.font.color.rgb = color
                    run.bold = True
        
        doc.add_paragraph()
        
        # Category summary
        doc.add_heading('2. Results by Category', level=1)
        
        by_cat = summary.get('by_category', {})
        cat_table = doc.add_table(rows=len(by_cat) + 1, cols=5)
        cat_table.style = 'Table Grid'
        
        cat_headers = ['Category', 'OK', 'Warning', 'Critical', 'Unknown']
        for i, h in enumerate(cat_headers):
            cat_table.rows[0].cells[i].text = h
            cat_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cat_table.rows[0].cells[i].paragraphs[0].runs:
                run.bold = True
        
        for row_idx, (cat_name, cat_data) in enumerate(by_cat.items(), start=1):
            cat_table.rows[row_idx].cells[0].text = cat_name
            cat_table.rows[row_idx].cells[1].text = str(cat_data.get('ok', 0))
            cat_table.rows[row_idx].cells[2].text = str(cat_data.get('warning', 0))
            cat_table.rows[row_idx].cells[3].text = str(cat_data.get('critical', 0))
            cat_table.rows[row_idx].cells[4].text = str(cat_data.get('unknown', 0))
            for cell in cat_table.rows[row_idx].cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # Detailed results
        doc.add_heading('3. Detailed Check Results', level=1)
        
        # Group by category
        categories = {'OS': [], 'Kubernetes': [], 'Services': []}
        for r in results:
            cat = r.get('Category', 'Unknown')
            if cat in categories:
                categories[cat].append(r)
        
        for cat_name, cat_results in categories.items():
            if not cat_results:
                continue
                
            doc.add_heading(f'3.{list(categories.keys()).index(cat_name)+1} {cat_name} Checks', level=2)
            
            # Table
            table = doc.add_table(rows=len(cat_results) + 1, cols=5)
            table.style = 'Table Grid'
            
            # Header
            headers = ['CheckID', 'Check Item', 'Status', 'Value', 'Message']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in table.rows[0].cells[i].paragraphs[0].runs:
                    run.bold = True
            
            # Data
            for row_idx, r in enumerate(cat_results, start=1):
                table.rows[row_idx].cells[0].text = r.get('CheckID', '')
                table.rows[row_idx].cells[1].text = r.get('CheckItem', '')
                
                status = r.get('Status', '')
                table.rows[row_idx].cells[2].text = status
                # Status colors
                status_cell = table.rows[row_idx].cells[2]
                for run in status_cell.paragraphs[0].runs:
                    if status == 'OK':
                        run.font.color.rgb = RGBColor(0, 128, 0)
                    elif status == 'Warning':
                        run.font.color.rgb = RGBColor(255, 165, 0)
                    elif status == 'Critical':
                        run.font.color.rgb = RGBColor(255, 0, 0)
                    run.bold = True
                
                # Truncate long values
                value = r.get('Value', '')
                if len(value) > 50:
                    value = value[:50] + "..."
                table.rows[row_idx].cells[3].text = value
                table.rows[row_idx].cells[4].text = r.get('Message', '')
            
            doc.add_paragraph()
        
        # Issues requiring action
        issues = [r for r in results if r.get('Status') in ['Warning', 'Critical']]
        if issues:
            doc.add_heading('4. Action Required Items', level=1)
            for issue in issues:
                para = doc.add_paragraph()
                status = issue.get('Status', '')
                icon = "⚠️" if status == 'Warning' else "❌"
                run = para.add_run(f"{icon} [{issue.get('CheckID')}] {issue.get('CheckItem')}")
                run.bold = True
                para.add_run(f"\n   - Status: {status}")
                para.add_run(f"\n   - Message: {issue.get('Message', '')}")
                para.add_run(f"\n   - Description: {issue.get('Description', '')}")
        
        # Signature section
        doc.add_paragraph()
        doc.add_paragraph()
        sign_para = doc.add_paragraph()
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        sign_para.add_run("Checked by: ________________")
        sign_para.add_run("\n\n")
        sign_para.add_run("Reviewed by: ________________")
        
        doc.save(filepath)
        return filepath


def generate_reports(results: List[Dict], summary: Dict, config: ReportConfig = None) -> Dict[str, str]:
    """Generate both CSV and DOCX reports"""
    generator = ReportGenerator(config)
    
    generated = {}
    
    # CSV
    csv_path = generator.generate_csv(results, summary)
    generated['csv'] = csv_path
    
    # DOCX
    if DOCX_AVAILABLE:
        docx_path = generator.generate_docx(results, summary)
        generated['docx'] = docx_path
    
    return generated


if __name__ == "__main__":
    # Test with dummy data
    test_results = [
        {'CheckID': 'OS-001', 'CheckItem': 'Disk Usage', 'Category': 'OS', 
         'Description': 'Test', 'Status': 'OK', 'Value': '45%', 'Threshold': '80%', 
         'Message': 'Within normal range', 'Timestamp': datetime.now().isoformat()},
    ]
    test_summary = {'total': 1, 'ok': 1, 'warning': 0, 'critical': 0, 'unknown': 0,
                    'by_category': {'OS': {'ok': 1, 'warning': 0, 'critical': 0, 'unknown': 0}}}
    
    config = ReportConfig(company_name="Test Corp", team_name="DevOps")
    paths = generate_reports(test_results, test_summary, config)
    print(f"Generated reports: {paths}")
